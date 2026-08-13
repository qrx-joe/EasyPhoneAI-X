from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "爸爸妈妈急救_项目说明书_参赛版.docx"

NAVY = "173B57"
BLUE = "2878A8"
TEAL = "2A8C82"
LIGHT_BLUE = "EAF4F8"
LIGHT_TEAL = "E9F5F2"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "68757F"
DARK = "20313D"
RED = "B33A3A"
GOLD = "C78B2A"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=DARK, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_text(p, text, bold=False, color=DARK, size=10.5):
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return r


def add_paragraph(doc, text="", *, bold_prefix=None, align=None, after=6, before=0, size=10.5, color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True, color=color, size=size)
        add_text(p, text[len(bold_prefix):], color=color, size=size)
    else:
        add_text(p, text, color=color, size=size)
    return p


def add_bullet(doc, text, *, level=0, bold_prefix=None, after=4):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True)
        add_text(p, text[len(bold_prefix):])
    else:
        add_text(p, text)
    return p


def add_number(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True)
        add_text(p, text[len(bold_prefix):])
    else:
        add_text(p, text)
    return p


def add_restarted_numbers(doc, items):
    """Add a real Word numbered list that restarts at 1."""
    style = doc.styles["List Number"]
    style_num_id = int(style._element.pPr.numPr.numId.val)
    numbering = doc.part.numbering_part.element
    source_num = next(n for n in numbering.num_lst if int(n.numId) == style_num_id)
    new_num = numbering.add_num(int(source_num.abstractNumId.val))
    new_num.add_lvlOverride(ilvl=0).add_startOverride(1)
    for text in items:
        p = doc.add_paragraph(style="List Number")
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = int(new_num.numId)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.2
        add_text(p, text)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    add_text(p, text, bold=True, color=NAVY if level == 1 else BLUE, size=16 if level == 1 else 12.5)
    return p


def add_callout(doc, title, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    add_text(p, title + "  ", bold=True, color=accent, size=10.5)
    add_text(p, text, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    for i, h in enumerate(headers):
        shade(table.rows[0].cells[i], header_fill)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, h, bold=True, color=WHITE, size=9.5)
    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            shade(cells[i], WHITE if r_i % 2 == 0 else LIGHT_GRAY)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if i == 0:
                add_text(p, value, bold=True, color=NAVY, size=9.3)
            else:
                add_text(p, value, size=9.3)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.76)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, before, after in ((1, 16, 16, 8), (2, 12.5, 11, 5), (3, 11, 8, 4)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    add_text(hp, "小有可为 2026｜普惠养老：银发数字生活", color=MID_GRAY, size=8.5)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    fp._p.append(field)


def build():
    doc = Document()
    configure_styles(doc)
    props = doc.core_properties
    props.title = "爸妈别急——银发数字生活安全副驾｜项目说明书"
    props.subject = "小有可为 2026 参赛项目说明"
    props.author = "爸妈别急项目组"
    props.keywords = "银发数字生活, 适老化, AI, 风险防护, 数字助老"

    # Cover: editorial cover pattern, intentionally restrained.
    for _ in range(3):
        add_paragraph(doc, "", after=22)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    add_text(p, "小有可为 2026｜普惠养老：银发数字生活", bold=True, color=TEAL, size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    add_text(p, "爸妈别急", bold=True, color=NAVY, size=30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    add_text(p, "银发数字生活安全副驾", bold=True, color=BLUE, size=17)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    add_text(p, "看懂当前一步，只给安全的下一步；遇到风险，先停下来。", color=MID_GRAY, size=11.5)
    add_callout(doc, "核心主张", "我们不替老人自动操作手机，而是帮助他们判断风险、完成可逆的小步骤，并在关键时刻顺利获得家人帮助。", LIGHT_TEAL, TEAL)
    add_paragraph(doc, "项目说明书｜参赛版", align=WD_ALIGN_PARAGRAPH.CENTER, color=MID_GRAY, size=10, before=32)
    add_paragraph(doc, "作品类型：应用类（Web 应用 / AI 辅助工具）", align=WD_ALIGN_PARAGRAPH.CENTER, color=MID_GRAY, size=9.5)
    add_paragraph(doc, "版本日期：2026 年 8 月", align=WD_ALIGN_PARAGRAPH.CENTER, color=MID_GRAY, size=9.5)

    page_break(doc)
    add_heading(doc, "一、项目概览", 1)
    add_paragraph(doc, "“爸妈别急”是一款面向银发用户的数字生活安全辅助应用。当老人面对看不懂的手机页面、复杂的设置路径，或退款、转账等潜在风险操作时，可通过文字、语音或截图描述当前情况。系统判断用户所处步骤与风险等级，一次只给出一个清晰、可逆且经过审核的下一步，并提供可观察的成功信号。")
    add_paragraph(doc, "对于转账、验证码、屏幕共享、陌生链接等高风险场景，产品停止普通指导，提示用户不要继续，并生成脱敏求助信息以便联系家人。对于信息不足、页面模糊或模型无法可靠判断的情况，系统进入“不确定”状态，不猜测、不强行继续。")
    add_callout(doc, "一句话介绍", "一个帮助老年人安全使用手机的 AI 副驾：低风险时一步一步教，高风险时立即停下并连接家人。")

    add_heading(doc, "二、问题与用户", 1)
    add_heading(doc, "2.1 用户面临的真实困难", 2)
    add_bullet(doc, "界面变化快：同一功能在不同手机、App 版本中的入口不同，记住教程也可能找不到。")
    add_bullet(doc, "说明不适老：常见教程一次给出多个步骤，字体小、术语多，老人难以确认自己是否做对。")
    add_bullet(doc, "安全风险高：退款、转账、验证码、屏幕共享等场景容易被诈骗话术利用。")
    add_bullet(doc, "求助成本高：老人难以准确描述当前页面，家人也需要反复追问才能理解问题。")
    add_heading(doc, "2.2 目标用户与使用者", 2)
    add_table(doc, ["角色", "核心需求", "产品价值"], [
        ("银发用户", "看懂当前页面；知道下一步；避免误操作", "大字、语音、一步式指导与风险停止"),
        ("子女 / 家属", "快速了解问题；减少重复远程沟通", "接收脱敏、结构化的求助信息"),
        ("社区助老人员", "提供一致、安全的数字技能辅导", "基于审核教程的标准化指导"),
    ], [1600, 3600, 4160])

    page_break(doc)
    add_heading(doc, "三、产品方案与核心流程", 1)
    add_paragraph(doc, "产品围绕“描述问题—识别风险—给出下一步—确认结果—必要时求助”构建闭环。系统首先使用确定性风险规则进行底线判断，再结合可选的视觉模型理解截图；最终输出只能落入四种受控决策：继续指导、停止操作、补充信息或暂不支持。")
    add_table(doc, ["阶段", "用户看到什么", "系统做什么"], [
        ("1. 描述问题", "输入文字、语音，或在同意后上传截图", "形成结构化观察信息"),
        ("2. 风险判断", "明确的安全提示与风险等级", "关键词规则与视觉结果取最高风险"),
        ("3. 下一步", "一次只显示一个动作及成功信号", "仅从人工审核的白名单教程选择"),
        ("4. 结果确认", "“完成了 / 没找到 / 需要帮助”等出口", "根据反馈继续、澄清或停止"),
        ("5. 家人协助", "可复制或分享的脱敏求助卡", "去除敏感信息并保留必要上下文"),
    ], [1500, 3800, 4060])
    add_heading(doc, "3.1 四种受控决策", 2)
    add_bullet(doc, "GUIDE（继续指导）：仅用于低风险且命中审核教程的场景。", bold_prefix="GUIDE（继续指导）：")
    add_bullet(doc, "STOP（停止操作）：风险达到 high / critical 时立即触发。", bold_prefix="STOP（停止操作）：")
    add_bullet(doc, "CLARIFY（补充信息）：截图模糊、信息冲突或模型失败时触发。", bold_prefix="CLARIFY（补充信息）：")
    add_bullet(doc, "UNSUPPORTED（暂不支持）：超出能力边界时明确告知，不伪造答案。", bold_prefix="UNSUPPORTED（暂不支持）：")
    add_heading(doc, "3.2 适老化交互原则", 2)
    add_bullet(doc, "一个页面只承担一个主要任务，减少同时出现的选择。")
    add_bullet(doc, "大字号、高对比度、短句表达，并提供语音输入与语音播报。")
    add_bullet(doc, "每一步说明“做什么”和“看到什么算成功”，降低不确定感。")
    add_bullet(doc, "危险操作不提供继续按钮，避免提示与行为相互矛盾。")

    page_break(doc)
    add_heading(doc, "四、核心功能", 1)
    add_table(doc, ["功能", "具体说明", "安全设计"], [
        ("多模态问题描述", "支持文字、语音与可选截图，降低描述手机页面的难度", "截图需明确同意；支持预览与本地手动遮挡"),
        ("一步式安全指导", "每次只给一个已审核动作，并显示成功信号", "动作必须可逆；模型自由文本不能直接驱动操作"),
        ("风险识别与拦截", "识别转账、验证码、屏幕共享、陌生渠道等危险信号", "规则命中高风险后，普通指导立即终止"),
        ("不确定状态", "信息不足或识别失败时，引导重新描述、补充信息或求助", "不确定不得进入普通指导，不采用 fail-open"),
        ("脱敏求助卡", "整理当前问题、风险提示和必要上下文，便于家人接手", "过滤验证码、账号等敏感内容和危险话术"),
        ("语音与无障碍", "语音输入、语音播报、大字号及高对比度界面", "设计令牌通过 WCAG 对比度自动检查"),
    ], [1700, 4200, 3460])
    add_heading(doc, "五、三条核心演示路径", 1)
    add_number(doc, "低风险｜微信没声音 / 字太小：系统识别问题后进入白名单教程，逐步提供一个可逆动作与成功信号。", bold_prefix="低风险｜")
    add_number(doc, "中风险｜电商退款：系统只允许经过审核的退款步骤；若页面出现付款、验证码或陌生渠道，立即升级风险。", bold_prefix="中风险｜")
    add_number(doc, "高风险｜转账 / 验证码 / 屏幕共享：系统直接停止指导，展示风险提醒并生成脱敏求助卡。", bold_prefix="高风险｜")
    add_callout(doc, "演示真实性", "三条路径使用同一套生产路由与安全规则，不设置隐藏的比赛演示分支。固定回放样例仅用于离线回归和无网络兜底，不被描述为真实模型调用。", LIGHT_GRAY, MID_GRAY)

    page_break(doc)
    add_heading(doc, "六、产品亮点", 1)
    add_heading(doc, "6.1 从“万能助手”转向“安全副驾”", 2)
    add_paragraph(doc, "传统助手追求尽快给出答案，本项目首先判断“现在是否应该继续”。面对银发用户与财产安全场景，停止、澄清和转交家人不是失败，而是产品能力的一部分。")
    add_heading(doc, "6.2 一次只给一个可逆步骤", 2)
    add_paragraph(doc, "系统不把整段教程一次性抛给用户，而是将复杂任务拆成小步骤。每一步都有明确动作和成功信号，用户可以随时停下，不会因为遗漏中间状态而越走越错。")
    add_heading(doc, "6.3 AI 只能升级风险，不能降低风险", 2)
    add_paragraph(doc, "确定性规则是安全底线。视觉模型可以发现额外风险并提高等级，但不能推翻规则已经识别出的高风险。该机制减少模型误判对用户安全的影响。")
    add_heading(doc, "6.4 不确定时不瞎猜", 2)
    add_paragraph(doc, "截图缺失、画面模糊、结果冲突、模型超时或结构无效时，系统进入 CLARIFY / UNKNOWN，而不是默认按低风险继续。")
    add_heading(doc, "6.5 从求助描述到家人接手", 2)
    add_paragraph(doc, "产品把老人难以表达的页面问题整理成结构化、脱敏的求助内容，使家人快速理解“正在做什么、哪里卡住、是否有风险”，缩短沟通链路。")
    add_callout(doc, "差异化价值", "不是“替老人点得更快”，而是“让老人每一步都更清楚、更可控，并在危险之前停下来”。", LIGHT_TEAL, TEAL)

    page_break(doc)
    add_heading(doc, "七、技术方案与特色", 1)
    add_heading(doc, "7.1 技术架构", 2)
    add_paragraph(doc, "项目采用 Next.js 16、React 19 与 TypeScript strict 构建，并使用六层分层架构隔离界面、业务编排、安全规则、跨层协议与外部 AI 服务。依赖方向保持单向，安全核心可在不启动浏览器和模型服务的情况下进行纯函数测试。")
    add_table(doc, ["层级", "主要职责", "典型内容"], [
        ("app / components", "页面、交互与适老化呈现", "首页、确认页、指导页、风险提醒页、语音组件"),
        ("application", "用例编排与决策流程", "观察屏幕、生成下一步、构建求助信息"),
        ("domain", "纯函数安全核心", "风险分类、教程匹配、路由保险丝、脱敏规则"),
        ("contracts", "跨边界数据协议", "UIObservation、GuidanceDecision、错误码"),
        ("infrastructure", "外部能力适配", "Qwen Vision 适配器、遥测接口"),
        ("lib", "无业务含义的通用能力", "语音、对比度与客户端工具"),
    ], [1900, 3400, 4060])
    add_heading(doc, "7.2 混合决策链", 2)
    add_paragraph(doc, "系统先执行关键词与场景规则，再按需调用视觉模型，最后通过风险取最大值策略合并结果。决策输出必须符合固定 Schema，模型自由文本不能直接控制页面跳转或点击。")
    add_callout(doc, "核心安全不变量", "高风险绝不进入 GUIDE；UNKNOWN 不得进入 GUIDE；视觉失败不得自动放行；每个指导步骤只包含一个审核动作与一个成功信号。", "FFF5E8", GOLD)
    add_heading(doc, "7.3 隐私与降级策略", 2)
    add_bullet(doc, "截图必须由用户主动选择，并在上传前进行预览和手动遮挡。")
    add_bullet(doc, "审计事件不保存用户原始描述文本，减少敏感信息暴露。")
    add_bullet(doc, "无 API Key 或模型不可用时，文字风险规则仍可运行；截图能力明确降级。")

    page_break(doc)
    add_heading(doc, "八、当前完成度与验证证据", 1)
    add_paragraph(doc, "本项目严格区分代码实现、本地测试、真实模型调用、线上运行与真实用户验证。以下状态基于当前仓库与本地验证记录，不将固定样例或代码推导结果描述为线上或真实设备证据。")
    add_table(doc, ["能力 / 验证项", "当前状态", "证据或说明"], [
        ("风险分类与最高风险合并", "已实现", "单元测试覆盖风险关键词与合并策略"),
        ("高风险路由保险丝", "已实现", "测试覆盖高风险不能进入普通教程"),
        ("求助卡与脱敏序列化", "已实现", "测试覆盖敏感信息与危险话术过滤"),
        ("白名单教程与退款场景", "已实现", "测试覆盖教程匹配和风险上限"),
        ("三条移动端主路径", "本地通过", "Playwright 覆盖低风险、中风险退款和高风险停止"),
        ("TypeScript / 生产构建", "本地通过", "strict typecheck 与 Next.js build 已通过"),
        ("真实 Qwen Vision 调用", "待验证", "适配器已实现，尚未使用真实 API Key 端到端测试"),
        ("线上部署", "未完成", "当前仅完成本地构建与运行验证"),
        ("真实老年用户测试", "未完成", "仍需开展可用性与理解度测试"),
    ], [2500, 1500, 5360])
    add_heading(doc, "8.1 当前边界", 2)
    add_bullet(doc, "不提供自动点击、远程控制或无障碍 Service，避免越权操作。")
    add_bullet(doc, "不建设家属账号、数据库和家属端收件箱；当前通过系统分享或复制求助信息完成交接。")
    add_bullet(doc, "多步任务当前以审核教程推进，完整 Allowed Action 状态机仍是后续工作。")
    add_bullet(doc, "截图遮挡已完成浏览器实现，但仍需真实手机触控与易用性验收。")

    page_break(doc)
    add_heading(doc, "九、社会价值与推广方向", 1)
    add_paragraph(doc, "“爸妈别急”聚焦老年人数字生活中的“最后一公里”：不仅帮助用户找到功能入口，也帮助他们判断什么时候不应该继续。项目可作为家庭数字助老工具，也可在社区数字课堂、银行与运营商适老服务、公益助老活动中提供标准化、安全化的辅助。")
    add_table(doc, ["应用方向", "可落地场景", "预期价值"], [
        ("家庭", "日常手机设置、购物、支付前确认", "降低反复沟通成本，增强老人的独立感"),
        ("社区", "数字技能课堂、志愿者辅导", "统一指导口径，优先识别风险"),
        ("公共服务", "医院、政务、交通 App 使用", "以一步式教程降低数字门槛"),
        ("金融与运营商", "反诈教育、适老服务入口", "在高风险节点提供前置提醒与人工转接"),
    ], [1600, 3800, 3960])
    add_heading(doc, "十、下一阶段计划", 1)
    add_restarted_numbers(doc, [
        "完成真实 Qwen Vision 端到端调用，记录响应质量、延迟、失败率与降级行为。",
        "邀请银发用户开展真实设备可用性测试，重点验证字体、语音、截图遮挡和风险提示的理解度。",
        "补齐多步任务状态机与更多人工审核任务包，同时保持动作白名单和可逆原则。",
        "完成线上部署、安全检查与隐私说明，在小范围社区或家庭中进行试点。",
    ])
    add_heading(doc, "十一、结语", 1)
    add_paragraph(doc, "面对不断变化的数字世界，老人需要的未必是一个替他们完成所有操作的“万能 AI”，而是一个知道什么时候指导、什么时候停下、什么时候请家人加入的安全伙伴。“爸妈别急”希望让每一次手机操作都更清楚、更可控，也让每一次求助都更容易被理解。")
    add_callout(doc, "项目愿景", "让每一位老人都能更有信心地使用数字服务，同时保留随时停下和寻求帮助的权利。", LIGHT_TEAL, TEAL)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
